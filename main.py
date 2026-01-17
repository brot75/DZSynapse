import customtkinter as ctk
import sys
from PIL import Image
import tkinter as tk
from tkinter import ttk, messagebox, filedialog
import database
import os
import shutil
import json
import uuid
from datetime import datetime
import tkcalendar
import docx
from docx.shared import Pt, Inches

# Load Wilaya/Commune Data
# Load Wilaya/Commune Data
WILAYAS = []
COMMUNES = {} # wilaya_id -> list of commune names

# NEOBRUTALISM STYLE
# STANDARD STYLE (Reverted to Default)
BUTTON_STYLE = {} 

# RED BUTTON (Delete/Logout)
BUTTON_STYLE_RED = {
    "fg_color": "#FF5252",
    "hover_color": "#FF1744"
}

# GREEN BUTTON (Add/New)
BUTTON_STYLE_GREEN = {
    "fg_color": "#00E676",
    "hover_color": "#00C853",
    "text_color": "black" # Fix visibility on bright green
}

# CHECKBOX STYLE (User Req: 'Flip' inspired colors)
# CHECKBOX STYLE (User Req: 'Flip' inspired colors)
CHECKBOX_STYLE = {
    "fg_color": "#0b76ef",      
    "border_color": ("gray60", "#e8e8eb"),  # Darker border in Light mode for visibility
    "border_width": 2,          
    "checkmark_color": "white", 
    "corner_radius": 4,         
    "hover_color": "#0b76ef",   
    "text_color": ("black", "white")       # Adaptive text color
}

def load_location_data():
    global WILAYAS, COMMUNES
    try:
        # Use relative path for portability (exe support)
        base_path = os.path.join(os.getcwd(), "Wilaya-Of-Algeria")
        if not os.path.exists(base_path):
             # Fallback for dev environment or alternative placement
             base_path = r"C:\xampp\htdocs\Wilaya-Of-Algeria" 
        with open(os.path.join(base_path, "Wilaya_Of_Algeria.json"), "r", encoding="utf-8") as f:
            wilaya_data = json.load(f)
            # Format: "1 - Adrar", etc.
            WILAYAS = [f"{w['id']} - {w['name']}" for w in wilaya_data]
            
        with open(os.path.join(base_path, "Commune_Of_Algeria.json"), "r", encoding="utf-8") as f:
            commune_data = json.load(f)
            for c in commune_data:
                w_id = c['wilaya_id']
                if w_id not in COMMUNES:
                    COMMUNES[w_id] = []
                COMMUNES[w_id].append(c['name'])
            # Sort communes
            for w_id in COMMUNES:
                COMMUNES[w_id].sort()
                
    except Exception as e:
        print(f"Error loading location data: {e}")
        WILAYAS = ["16 - Alger"] # Fallback
        COMMUNES = {"16": ["Alger Centre", "Bab El Oued"]}

load_location_data()

# Configuration
ctk.set_appearance_mode("Dark")
ctk.set_default_color_theme("dark-blue")

KNOWLEDGE_BASE = {
    "Tête et Cou": ["Kystes thyréoglosses", "Fentes branchiales", "Adénopathies"],
    "Thorax": ["Malformations pulmonaires (CPAM)", "Hernie diaphragmatique (Bochdalek/Morgagni)", "Pectus excavatum/carinatum"],
    "Digestif": ["Atrésie œsophage", "Sténose du pylore", "Maladie de Hirschsprung", "Malformations anorectales", "Invagination"],
    "Uro-Génital": ["Cryptorchidie", "Hypospadias", "Tumeur de Wilms", "Torsion testiculaire"],
    "Paroi Abdominale": ["Hernie inguinale", "Ombilic", "Gastroschisis", "Omphalocèle"],
    "Oncologie": ["Neuroblastome", "Tératome sacro-coccygien"],
    "Trauma/Brûlures": ["Gestion des voies aériennes", "Brûlures (Liquides)", "Traumas fermés"]
}





class LicenseFrame(ctk.CTkFrame):
    def __init__(self, master, callback):
        super().__init__(master)
        self.callback = callback
        
        self.pack(fill="both", expand=True)
        
        self.center_frame = ctk.CTkFrame(self)
        self.center_frame.place(relx=0.5, rely=0.5, anchor="center")
        
        ctk.CTkLabel(self.center_frame, text="Activation Requise", font=("Roboto", 24, "bold")).pack(pady=20)
        ctk.CTkLabel(self.center_frame, text="Veuillez entrer votre clé de licence.", font=("Roboto", 12)).pack(pady=(0, 20))
        
        self.key_entry = ctk.CTkEntry(self.center_frame, placeholder_text="Clé de licence", show="*")
        self.key_entry.pack(pady=10, padx=20)
        
        self.btn = ctk.CTkButton(self.center_frame, text="Activer", command=self.activate, **BUTTON_STYLE_GREEN)
        self.btn.pack(pady=20)
        
    def activate(self):
        key = self.key_entry.get()
        if key == "khaled2026":
            self.callback()
        else:
            messagebox.showerror("Erreur", "Clé de licence invalide.")

class LoginFrame(ctk.CTkFrame):
    def __init__(self, master, login_callback):
        super().__init__(master)
        self.login_callback = login_callback
        
        self.pack(fill="both", expand=True)
        
        # White background for container to contrast with Gray Inputs
        self.center_frame = ctk.CTkFrame(self)
        self.center_frame.place(relx=0.5, rely=0.5, anchor="center")
        
        ctk.CTkLabel(self.center_frame, text="DZSynapse", font=("Roboto", 24, "bold")).pack(pady=20)
        
        self.username_entry = ctk.CTkEntry(self.center_frame, placeholder_text="Nom d'utilisateur")
        self.username_entry.pack(pady=10, padx=20)
        
        self.password_entry = ctk.CTkEntry(self.center_frame, placeholder_text="Mot de passe", show="*")
        self.password_entry.pack(pady=10, padx=20)
        
        self.remember_var = ctk.BooleanVar(value=False)
        self.remember_cb = ctk.CTkCheckBox(self.center_frame, text="Rester connecté", variable=self.remember_var, **CHECKBOX_STYLE)
        self.remember_cb.pack(pady=5)
        
        self.login_button = ctk.CTkButton(self.center_frame, text="Connexion", command=self.login_event, **BUTTON_STYLE)
        self.login_button.pack(pady=20)
        
    def login_event(self):
        username = self.username_entry.get()
        password = self.password_entry.get()
        
        user = database.verify_login(username, password)
        if user:
            if self.remember_var.get():
                token = database.create_session(username)
                try:
                    with open("session.json", "w") as f:
                        json.dump({"username": username, "token": token}, f)
                except Exception as e:
                    print(f"Failed to save session: {e}")
            self.login_callback(user)
        else:
            messagebox.showerror("Erreur", "Nom d'utilisateur ou mot de passe incorrect")

class KnowledgeBaseFrame(ctk.CTkFrame):
    def __init__(self, master):
        super().__init__(master)
        
        self.label = ctk.CTkLabel(self, text="Référentiel Procédures (Royal College Canada)", font=("Roboto", 20, "bold"))
        self.label.pack(pady=20, padx=20, anchor="w")
        
        # Search bar
        self.search_var = ctk.StringVar()
        self.search_var.trace("w", self.filter_tree)
        self.search_entry = ctk.CTkEntry(self, placeholder_text="Rechercher une compétence...", textvariable=self.search_var)
        self.search_entry.pack(fill="x", padx=20, pady=(0, 20))

        # Use ttk Treeview for the hierarchy
        self.tree_frame = ctk.CTkFrame(self)
        self.tree_frame.pack(fill="both", expand=True, padx=20, pady=20)
        
        self.style = ttk.Style()
        self.style.theme_use("default")
        self.style.configure("Treeview", 
                             background="#2b2b2b", 
                             foreground="white", 
                             fieldbackground="#2b2b2b", 
                             font=("Roboto", 12))
        self.style.map("Treeview", background=[("selected", "#1f538d")])
        
        self.tree = ttk.Treeview(self.tree_frame, columns=("item"), show="tree")
        self.tree.pack(fill="both", expand=True, side="left")
        
        self.scrollbar = ttk.Scrollbar(self.tree_frame, orient="vertical", command=self.tree.yview)
        self.scrollbar.pack(side="right", fill="y")
        self.tree.configure(yscrollcommand=self.scrollbar.set)
        
        self.populate_tree()

    def populate_tree(self):
        # Clear existing
        for item in self.tree.get_children():
            self.tree.delete(item)
            
        for category, items in KNOWLEDGE_BASE.items():
            cat_id = self.tree.insert("", "end", text=category, open=True)
            for item in items:
                self.tree.insert(cat_id, "end", text=item)
                
    def filter_tree(self, *args):
        query = self.search_var.get().lower()
        if not query:
            self.populate_tree()
            return
            
        for item in self.tree.get_children():
            self.tree.delete(item)
            
        for category, items in KNOWLEDGE_BASE.items():
            # Check if category matches or any item matches
            matching_items = [i for i in items if query in i.lower()]
            category_match = query in category.lower()
            
            if category_match or matching_items:
                cat_id = self.tree.insert("", "end", text=category, open=True)
                # If category matches, show all items? Or just show the category? 
                # Let's show all items if category matches, OR only matching items if not.
                items_to_show = items if category_match else matching_items
                for item in items_to_show:
                    self.tree.insert(cat_id, "end", text=item)


class CustomDatePicker(ctk.CTkFrame):
    def __init__(self, master, width=140, placeholder_text="", command=None, **kwargs):
        super().__init__(master, width=width, height=32, fg_color="transparent")
        self.command = command
        self.date_str = ctk.StringVar()
        
        self.entry = ctk.CTkEntry(self, textvariable=self.date_str, width=width-40, placeholder_text=placeholder_text)
        self.entry.pack(side="left", padx=(0, 5))
        
        self.btn = ctk.CTkButton(self, text="📅", width=30, command=self.open_calendar, fg_color="#1f538d")
        self.btn.pack(side="left")
        
    def open_calendar(self):
        # Position logic
        try:
            x = self.entry.winfo_rootx()
            y = self.entry.winfo_rooty() + self.entry.winfo_height() + 5
        except:
            x = 500
            y = 300
            
        self.top = ctk.CTkToplevel(self)
        self.top.geometry(f"+{x}+{y}")
        self.top.title("Choisir Date")
        self.top.attributes("-topmost", True)
        self.top.grab_set() # Modal
        
        self.cal = tkcalendar.Calendar(self.top, selectmode='day', date_pattern='dd/mm/yyyy',
                                       background='darkblue', foreground='white', borderwidth=0)
        self.cal.pack(padx=10, pady=10)
        
        # Try to set initial date from entry
        try:
            current = self.date_str.get()
            if current:
                 d = datetime.strptime(current, "%d/%m/%Y")
                 self.cal.selection_set(d)
        except: pass

        ctk.CTkButton(self.top, text="Valider", command=self.select_date, width=100).pack(pady=10)
        
    def select_date(self):
        date = self.cal.get_date()
        self.date_str.set(date)
        if self.command:
            self.command(date)
        self.top.destroy()
        
    def get_date(self):
        try:
            dt = datetime.strptime(self.date_str.get(), "%d/%m/%Y")
            return dt.date() # Return date object like DateEntry
        except:
            return None
            
    def set_date(self, date_obj):
        if isinstance(date_obj, str):
             self.date_str.set(date_obj)
        elif hasattr(date_obj, 'strftime'):
             self.date_str.set(date_obj.strftime("%d/%m/%Y"))
             
    def get(self):
        return self.date_str.get()
        
    def delete(self, first, last=None):
        # Mimic Entry.delete
        self.entry.delete(first, last)
    
    def insert(self, index, string):
        self.entry.insert(index, string)
        
    def bind(self, sequence=None, command=None, add=None):
        pass
        
    def configure(self, **kwargs):
        if 'state' in kwargs:
            self.entry.configure(state=kwargs['state'])
            self.btn.configure(state=kwargs['state'])


class PatientManagementFrame(ctk.CTkFrame):
    def __init__(self, master, role):
        super().__init__(master)
        self.role = role
        self.current_patient = None
        
        # Split layout: List (Left) and Details (Right)
        self.grid_columnconfigure(1, weight=1)
        self.grid_rowconfigure(0, weight=1)
        
        # Left Panel: Search and List
        self.left_panel = ctk.CTkFrame(self, width=250)
        self.left_panel.grid(row=0, column=0, sticky="nswe", padx=10, pady=10)
        
        self.search_entry = ctk.CTkEntry(self.left_panel, placeholder_text="Nom, Prénom, NSS, Tél...")
        self.search_entry.pack(pady=10, padx=10, fill="x")
        self.search_entry.bind("<Return>", self.search_patients)
        self.search_entry.bind("<KeyRelease>", self.search_patients)
        
        self.search_btn = ctk.CTkButton(self.left_panel, text="Rechercher", command=self.search_patients, **BUTTON_STYLE)
        self.search_btn.pack(pady=5, padx=10)
        
        self.add_btn = ctk.CTkButton(self.left_panel, text="Nouveau Patient", command=self.clear_form, **BUTTON_STYLE_GREEN)
        if self.role == 'doctor':
            self.add_btn.pack(pady=10, padx=10)
        
        self.patient_list_frame = ctk.CTkScrollableFrame(self.left_panel)
        self.patient_list_frame.pack(fill="both", expand=True, padx=5, pady=5)
        
        # Right Panel: Details with Tabs
        self.right_panel = ctk.CTkFrame(self)
        self.right_panel.grid(row=0, column=1, sticky="nswe", padx=10, pady=10)

        if self.role == 'operator':
            ctk.CTkLabel(self.right_panel, text="MODE LECTURE SEULE", text_color="red", font=("Roboto", 16, "bold")).pack(pady=5)

        self.tabview = ctk.CTkTabview(self.right_panel)
        self.tabview.pack(fill="both", expand=True, padx=10, pady=10)
        
        self.tab_info = self.tabview.add("Informations")
        self.tab_interventions = self.tabview.add("Interventions")
        self.tab_documents = self.tabview.add("Documents")
        
        # --- Tab 1: Informations (Form) ---
        self.form_frame = ctk.CTkScrollableFrame(self.tab_info)
        self.form_frame.pack(fill="both", expand=True, padx=5, pady=5)
        
        self.widgets = {}

        # 0. Motif de Consultation
        ctk.CTkLabel(self.form_frame, text="Motif de Consultation:", font=("Roboto", 12, "bold")).pack(anchor="w", padx=10, pady=(10, 0))
        self.widgets['motif_consultation'] = ctk.CTkEntry(self.form_frame, placeholder_text="Raison de la visite...")
        self.widgets['motif_consultation'].pack(fill="x", padx=10, pady=5)

        # 1. Identity
        id_frame = ctk.CTkFrame(self.form_frame, fg_color="transparent")
        id_frame.pack(fill="x", padx=5, pady=5)
        
        # Row 1: Nom, Prenom, NSS
        ctk.CTkLabel(id_frame, text="Nom:").grid(row=0, column=0, padx=5, sticky="e")
        self.widgets['nom'] = ctk.CTkEntry(id_frame, width=150); self.widgets['nom'].grid(row=0, column=1, padx=5)
        
        ctk.CTkLabel(id_frame, text="Prénom:").grid(row=0, column=2, padx=5, sticky="e")
        self.widgets['prenom'] = ctk.CTkEntry(id_frame, width=150); self.widgets['prenom'].grid(row=0, column=3, padx=5)
        
        ctk.CTkLabel(id_frame, text="NSS:").grid(row=0, column=4, padx=5, sticky="e")
        self.widgets['nss'] = ctk.CTkEntry(id_frame, width=150, placeholder_text="Numéro Chifa"); self.widgets['nss'].grid(row=0, column=5, padx=5)

        # Row 2: Date, Sexe
        ctk.CTkLabel(id_frame, text="Né(e) le:").grid(row=1, column=0, padx=5, pady=5, sticky="n")
        
        date_box = ctk.CTkFrame(id_frame, fg_color="transparent")
        date_box.grid(row=1, column=1, padx=5, pady=5, sticky="w")
        
        self.widgets['date_naissance'] = CustomDatePicker(date_box, width=150, placeholder_text="JJ/MM/AAAA", command=self.on_date_change)
        self.widgets['date_naissance'].pack()
        # BIGGER AGE LABEL
        self.age_label = ctk.CTkLabel(date_box, text="", font=("Roboto", 14, "bold"), text_color="cyan")
        self.age_label.pack()

        ctk.CTkLabel(id_frame, text="Sexe:").grid(row=1, column=2, padx=5, sticky="n", pady=5)
        self.widgets['sexe'] = ctk.CTkOptionMenu(id_frame, values=["Masculin", "Féminin"], width=150)
        self.widgets['sexe'].grid(row=1, column=3, padx=5, sticky="n", pady=5)

        # Row 3: Wilaya, Commune
        ctk.CTkLabel(id_frame, text="Wilaya:").grid(row=2, column=0, padx=5, pady=5, sticky="e")
        self.widgets['wilaya'] = ctk.CTkOptionMenu(id_frame, values=WILAYAS, width=150, command=self.update_communes)
        self.widgets['wilaya'].grid(row=2, column=1, padx=5, pady=5)
        
        ctk.CTkLabel(id_frame, text="Commune:").grid(row=2, column=2, padx=5, sticky="e")
        self.widgets['commune'] = ctk.CTkOptionMenu(id_frame, values=["Selectionnez"], width=150)
        self.widgets['commune'].grid(row=2, column=3, padx=5)

        # 2. Vitals (Poids, Taille, BMI)
        vitals_frame = ctk.CTkFrame(self.form_frame)
        vitals_frame.pack(fill="x", padx=10, pady=10)
        ctk.CTkLabel(vitals_frame, text="Constantes Vitales", font=("Roboto", 12, "bold")).pack(anchor="w", padx=10, pady=5)
        
        vf_inner = ctk.CTkFrame(vitals_frame, fg_color="transparent")
        vf_inner.pack(fill="x", padx=5, pady=5)
        
        ctk.CTkLabel(vf_inner, text="Poids (kg):").pack(side="left", padx=5)
        self.widgets['poids'] = ctk.CTkEntry(vf_inner, width=80)
        self.widgets['poids'].pack(side="left", padx=5)
        self.widgets['poids'].bind("<KeyRelease>", self.calculate_bmi)
        
        ctk.CTkLabel(vf_inner, text="Taille (cm):").pack(side="left", padx=5)
        self.widgets['taille'] = ctk.CTkEntry(vf_inner, width=80)
        self.widgets['taille'].pack(side="left", padx=5)
        self.widgets['taille'].bind("<KeyRelease>", self.calculate_bmi)
        
        self.bmi_label = ctk.CTkLabel(vf_inner, text="IMC: --", font=("Roboto", 12, "bold"), text_color="cyan")
        self.bmi_label.pack(side="left", padx=20)

        # 3. Volet Pédiatrique / Informations Naissance
        self.pediatric_frame = ctk.CTkFrame(self.form_frame, border_color="lightblue", border_width=2)
        
        ctk.CTkLabel(self.pediatric_frame, text="Informations de Naissance (Enfants)", font=("Roboto", 12, "bold"), text_color="lightblue").pack(pady=5)
        pf_inner = ctk.CTkFrame(self.pediatric_frame, fg_color="transparent")
        pf_inner.pack(fill="x", padx=5, pady=5)
        
        ctk.CTkLabel(pf_inner, text="Terme (SA):").grid(row=0, column=0, padx=5)
        self.widgets['terme_grossesse'] = ctk.CTkEntry(pf_inner, width=100)
        self.widgets['terme_grossesse'].grid(row=0, column=1, padx=5)
        
        ctk.CTkLabel(pf_inner, text="Poids Naissance (g):").grid(row=0, column=2, padx=5)
        self.widgets['poids_naissance'] = ctk.CTkEntry(pf_inner, width=100)
        self.widgets['poids_naissance'].grid(row=0, column=3, padx=5)
        
        ctk.CTkLabel(pf_inner, text="APGAR (1/5min):").grid(row=0, column=4, padx=5)
        self.widgets['apgar'] = ctk.CTkEntry(pf_inner, width=100)
        self.widgets['apgar'].grid(row=0, column=5, padx=5)
        
        ctk.CTkLabel(pf_inner, text="Accouchement:").grid(row=1, column=0, padx=5, pady=5)
        # FIX: "Vaginale" -> "Normale"
        self.widgets['mode_accouchement'] = ctk.CTkOptionMenu(pf_inner, values=["Normale", "Césarienne (Prog)", "Césarienne (Urg)"], width=150)
        self.widgets['mode_accouchement'].grid(row=1, column=1, padx=5, pady=5)

        # 4. Medical Flags
        flags_frame = ctk.CTkFrame(self.form_frame)
        flags_frame.pack(fill="x", padx=10, pady=10)
        
        ctk.CTkLabel(flags_frame, text="Groupe Sanguin:").grid(row=0, column=0, padx=10, pady=5, sticky="e")
        self.widgets['groupe_sanguin'] = ctk.CTkOptionMenu(flags_frame, values=["A+", "A-", "B+", "B-", "AB+", "AB-", "O+", "O-"], width=100)
        self.widgets['groupe_sanguin'].grid(row=0, column=1, padx=5, sticky="w")
        
        ctk.CTkLabel(flags_frame, text="Allergies (Majeures):", text_color="red", font=("Roboto", 12, "bold")).grid(row=0, column=2, padx=10, sticky="e")
        self.widgets['allergies'] = ctk.CTkComboBox(flags_frame, values=["Aucune", "Pénicilline", "Aspirine", "Arachide", "Lactose", "Latex"], width=200, text_color="red")
        self.widgets['allergies'].grid(row=0, column=3, padx=5, sticky="w")
        
        # Vaccins Multi-Select (Moved to Right of Allergies)
        ctk.CTkLabel(flags_frame, text="Vaccins:").grid(row=0, column=4, padx=10, sticky="e")
        self.widgets['vaccins'] = ctk.CTkTextbox(flags_frame, width=200, height=30)
        self.widgets['vaccins'].grid(row=0, column=5, padx=5, sticky="w")
        self.vaccin_btn = ctk.CTkButton(flags_frame, text="+", width=30, command=self.open_vaccine_dialog, **BUTTON_STYLE)
        self.vaccin_btn.grid(row=0, column=6, padx=5, sticky="w")

        # 5. History (Antécédents)
        history_frame = ctk.CTkFrame(self.form_frame)
        history_frame.pack(fill="x", padx=10, pady=10)
        ctk.CTkLabel(history_frame, text="Antécédents & Traitements", font=("Roboto", 12, "bold")).pack(anchor="w", padx=10, pady=5)
        
        h_grid = ctk.CTkFrame(history_frame, fg_color="transparent")
        h_grid.pack(fill="x", padx=5)
        
        # Helper lists for history
        surg_opts = ["Circoncision", "Hernie Inguinale", "Hypospadias", "Ectopie Testiculaire", "Amygdalectomie", "Végétations", "Autre"]
        fam_opts = ["Atopie/Allergie", "Asthme", "Diabète (Type 1)", "Epilepsie", "Consanguinité", "Autre"]
        
        # Médicaux (Text Only)
        ctk.CTkLabel(h_grid, text="Médicaux").grid(row=0, column=0, sticky="w", padx=5)
        self.widgets['ant_medicaux'] = ctk.CTkTextbox(h_grid, width=250, height=80)
        self.widgets['ant_medicaux'].grid(row=1, column=0, padx=5, pady=5)
        
        # Chirurgicaux (Dropdown + Text)
        ctk.CTkLabel(h_grid, text="Chirurgicaux").grid(row=0, column=1, sticky="w", padx=5)
        self.chir_combo = ctk.CTkComboBox(h_grid, values=surg_opts, width=250, command=lambda v: self.append_history('ant_chirurgicaux', v))
        self.chir_combo.grid(row=1, column=1, padx=5, pady=(0, 85), sticky="n") # Overlap? No, simple grid.
        self.widgets['ant_chirurgicaux'] = ctk.CTkTextbox(h_grid, width=250, height=60)
        self.widgets['ant_chirurgicaux'].grid(row=1, column=1, padx=5, pady=(30, 0)) # Push down
        
        # Familiaux (Dropdown + Text)
        ctk.CTkLabel(h_grid, text="Familiaux").grid(row=0, column=2, sticky="w", padx=5)
        self.fam_combo = ctk.CTkComboBox(h_grid, values=fam_opts, width=250, command=lambda v: self.append_history('ant_familiaux', v))
        self.fam_combo.grid(row=1, column=2, padx=5, pady=(0, 85), sticky="n")
        self.widgets['ant_familiaux'] = ctk.CTkTextbox(h_grid, width=250, height=60)
        self.widgets['ant_familiaux'].grid(row=1, column=2, padx=5, pady=(30, 0))
        
        # Traitements
        ctk.CTkLabel(h_grid, text="Traitements en Cours").grid(row=2, column=0, sticky="w", padx=5, pady=(10,0))
        self.widgets['traitements_actuels'] = ctk.CTkTextbox(h_grid, width=520, height=60)
        self.widgets['traitements_actuels'].grid(row=3, column=0, columnspan=2, padx=5, pady=5, sticky="ew")

        # Tuteur & Phone (Bottom)
        t_frame = ctk.CTkFrame(self.form_frame, fg_color="transparent")
        t_frame.pack(fill="x", padx=10, pady=10)
        
        ctk.CTkLabel(t_frame, text="Contact / Tuteur:", font=("Roboto", 12, "bold")).pack(anchor="w", padx=10)
        
        t_inner = ctk.CTkFrame(t_frame, fg_color="transparent")
        t_inner.pack(fill="x", padx=5, pady=5)
        
        ctk.CTkLabel(t_inner, text="N° Téléphone (+213):").pack(side="left", padx=5)
        self.widgets['telephone'] = ctk.CTkEntry(t_inner, width=150, placeholder_text="550 35 23 33")
        self.widgets['telephone'].pack(side="left", padx=5)

        ctk.CTkFrame(t_inner, width=20, height=1, fg_color="transparent").pack(side="left") # Spacer

        self.tuteur_lien_var = ctk.StringVar(value="Père")
        ctk.CTkRadioButton(t_inner, text="Père", variable=self.tuteur_lien_var, value="Père").pack(side="left", padx=5)
        ctk.CTkRadioButton(t_inner, text="Mère", variable=self.tuteur_lien_var, value="Mère").pack(side="left", padx=5)
        ctk.CTkRadioButton(t_inner, text="Autre", variable=self.tuteur_lien_var, value="Autre").pack(side="left", padx=5)
        
        self.widgets['tuteur_nom'] = ctk.CTkEntry(t_inner, placeholder_text="Nom Prénom du tuteur", width=200)
        self.widgets['tuteur_nom'].pack(side="left", padx=10)

        # Disable for Operator
        if self.role == 'operator':
            for w in self.widgets.values():
                try: w.configure(state="disabled")
                except: pass
            try: self.vaccin_btn.configure(state="disabled")
            except: pass

        # Action Buttons (Fixed at bottom of tab)
        self.btn_frame = ctk.CTkFrame(self.tab_info, fg_color="transparent")
        self.btn_frame.pack(pady=10)
        
        if self.role == 'doctor':
            self.save_btn = ctk.CTkButton(self.btn_frame, text="Enregistrer Dossier", command=self.save_patient, width=200, height=40, **BUTTON_STYLE)
            self.save_btn.pack(side="left", padx=10)
            
            self.export_btn = ctk.CTkButton(self.btn_frame, text="Exporter (DOCX)", command=self.export_patient_docx, width=150, **BUTTON_STYLE)
            self.export_btn.pack(side="left", padx=10)
            
            self.add_btn_bottom = ctk.CTkButton(self.btn_frame, text="Nouveau", command=self.clear_form, width=100, **BUTTON_STYLE_GREEN)
            self.add_btn_bottom.pack(side="left", padx=10)

            self.delete_btn = ctk.CTkButton(self.btn_frame, text="Supprimer", command=self.delete_patient, **BUTTON_STYLE_RED)
            self.delete_btn.pack(side="left", padx=10)
            
        # --- Tab 2 & 3 Init ---
        # Styling Tabs to match Buttons (Cyan/Black)
        # Styling Tabs to match Buttons (Cyan/Black)
        self.tab_info.configure(fg_color="transparent") # Content bg
        self.tabview.configure(segmented_button_fg_color="gray90",
                                   segmented_button_selected_color="#4CC9F0",
                                   segmented_button_selected_hover_color="#4895EF",
                                   segmented_button_unselected_color="gray90",
                                   segmented_button_unselected_hover_color="gray80",
                                   text_color="black") # Active text color assumption
        
        self.interv_list = ctk.CTkScrollableFrame(self.tab_interventions)
        self.interv_list.pack(fill="both", expand=True, padx=10, pady=10)
        if self.role == 'doctor':
             ctk.CTkButton(self.tab_interventions, text="Ajouter Intervention", command=self.add_intervention_dialog, **BUTTON_STYLE).pack(pady=10)
             
        self.doc_list = ctk.CTkScrollableFrame(self.tab_documents)
        self.doc_list.pack(fill="both", expand=True, padx=10, pady=10)
        if self.role == 'doctor':
            ctk.CTkButton(self.tab_documents, text="Ajouter Document", command=self.add_document_dialog, **BUTTON_STYLE).pack(pady=10)

        # Initial Load
        self.load_patients_list()
    
    def append_history(self, widget_name, value):
        if not value: return
        widget = self.widgets[widget_name]
        current_text = widget.get("1.0", "end-1c")
        if current_text.strip():
            widget.insert("end", f"\n- {value}")
        else:
            widget.insert("1.0", f"- {value}")
        
    def open_vaccine_dialog(self):
        dialog = ctk.CTkToplevel(self)
        dialog.title("Vaccins (Multi-sélection)")
        dialog.geometry("300x500")
        dialog.grab_set()
        
        # Read current lines
        current_text = self.widgets['vaccins'].get("1.0", "end-1c")
        current_list = [line.replace("- ", "").strip() for line in current_text.split("\n") if line.strip()]

        vars = {}
        options = ["BCG", "Hépatite B", "DTCoq", "Polio", "ROR", "Pneumo", "Méningocoque C", "Rotavirus"]
        
        scroll = ctk.CTkScrollableFrame(dialog)
        scroll.pack(fill="both", expand=True, padx=10, pady=10)
        
        for opt in options:
            var = ctk.BooleanVar(value=opt in current_list)
            ctk.CTkCheckBox(scroll, text=opt, variable=var, **CHECKBOX_STYLE).pack(anchor="w", padx=20, pady=5)
            vars[opt] = var
            
        def apply():
            selected = [k for k, v in vars.items() if v.get()]
            self.widgets['vaccins'].delete("1.0", "end")
            # Format as bullet list
            formatted = "\n".join([f"- {s}" for s in selected])
            self.widgets['vaccins'].insert("1.0", formatted)
            dialog.destroy()
            
        ctk.CTkButton(dialog, text="Valider", command=apply, **BUTTON_STYLE).pack(pady=20)

    def on_date_change(self, event):
        self.check_pediatric()
        
    def calculate_bmi(self, event=None):
        try:
            p = float(self.widgets['poids'].get())
            t = float(self.widgets['taille'].get()) / 100 # cm to m
            if t > 0:
                bmi = p / (t * t)
                self.bmi_label.configure(text=f"IMC: {bmi:.1f}")
            else:
                self.bmi_label.configure(text="IMC: --")
        except:
             self.bmi_label.configure(text="IMC: --")

    def check_pediatric(self):
        try:
            # Default: Show if undefined or child
            show_section = True 
            age_text = ""
            is_child = True
            
            dob = self.widgets['date_naissance'].get_date()
            if hasattr(dob, 'date'): dob = dob.date() 
            
            if dob:
                today = datetime.today().date()
                # Precise age calc
                years = today.year - dob.year - ((today.month, today.day) < (dob.month, dob.day))
                
                # Label Update
                if years < 2:
                    delta = today - dob
                    months = int(delta.days / 30.44)
                    weeks = int(delta.days / 7)
                    age_text = f"{months} mois ({weeks} semaines)"
                else:
                    age_text = f"{years} ans"
                
                # Logic: Show if < 16 years
                if years >= 16:
                    show_section = False
                    is_child = False
            
            self.age_label.configure(text=age_text)
            
            # --- Pediatric Surgery Exclusive Options ---
            ped_surg = ["Circoncision", "Hernie Inguinale", "Amygdalectomie", "Végétations", "Appendicectomie", "Fracture", "Invagination", "Hypospadias", "Ectopie Testiculaire", "Autre"]
            ped_fam = ["Atopie/Allergie", "Asthme", "Diabète (Type 1)", "Epilepsie", "Consanguinité", "Surdité", "Maladie Héréditaire", "Autre"]
            
            self.chir_combo.configure(values=ped_surg)
            self.fam_combo.configure(values=ped_fam)
            # -----------------------------------

            if show_section:
                self.pediatric_frame.pack(after=self.form_frame.winfo_children()[3], fill="x", padx=10, pady=10) 
            else:
                self.pediatric_frame.pack_forget()
                
        except Exception as e:
            print(f"Age Calc Error: {e}")

    def update_communes(self, selected_wilaya):
        if not selected_wilaya: return
        w_id = selected_wilaya.split(" - ")[0]
        communes = COMMUNES.get(w_id, [])
        self.widgets['commune'].configure(values=communes)
        if communes:
            self.widgets['commune'].set(communes[0])
        else:
            self.widgets['commune'].set("Selectionnez")

    def search_patients(self, event=None):
        query = self.search_entry.get()
        self.load_patients_list(query)

    def load_patients_list(self, query=""):
        for widget in self.patient_list_frame.winfo_children():
            widget.destroy()
            
        patients = database.get_all_patients(query)
        for p in patients:
            btn = ctk.CTkButton(self.patient_list_frame, text=f"{p['nom']} {p['prenom']}", 
                                fg_color="transparent", border_width=1, text_color=("gray10", "gray90"),
                                command=lambda pid=p['id']: self.load_patient(pid))
            btn.pack(fill="x", pady=2)

    def load_patient(self, pid):
        try:
            patient_row = database.get_patient(pid)
            if not patient_row: return
                
            # Convert to dict to allow .get() method
            patient = dict(patient_row)
            self.current_patient = patient
            
            # Helper to safely set ENTRY/COMBOBOX
            def safe_set(widget_name, value):
                try:
                    w = self.widgets.get(widget_name)
                    if not w: return
                    if isinstance(w, ctk.CTkEntry):
                        w.delete(0, "end")
                        w.insert(0, str(value) if value is not None else "")
                    elif isinstance(w, (ctk.CTkOptionMenu, ctk.CTkComboBox)):
                        w.set(str(value) if value is not None else "")
                    elif isinstance(w, ctk.CTkTextbox):
                        w.delete("1.0", "end")
                        w.insert("1.0", str(value) if value is not None else "")
                except Exception as e:
                    print(f"Error setting {widget_name}: {e}")

            # Map DB fields to widgets
            field_map = {
                'nom': patient.get('nom'), 'prenom': patient.get('prenom'), 
                'nss': patient.get('nss'), 'telephone': patient.get('telephone'), 
                'poids': patient.get('poids'), 'taille': patient.get('taille'), 
                'terme_grossesse': patient.get('terme_grossesse'), 
                'poids_naissance': patient.get('poids_naissance'), 
                'apgar': patient.get('apgar'), 'tuteur_nom': patient.get('tuteur_nom'), 
                'motif_consultation': patient.get('motif_consultation'),
                
                'sexe': patient.get('sexe'), 'wilaya': patient.get('wilaya'), 
                'commune': patient.get('commune'), 'groupe_sanguin': patient.get('groupe_sanguin'),
                'mode_accouchement': "Normale" if patient.get('mode_accouchement') == "Vaginale" else patient.get('mode_accouchement'), 
                'allergies': patient.get('allergies'), 'vaccins': patient.get('vaccins'),
                
                'ant_medicaux': patient.get('ant_medicaux'), 
                'ant_chirurgicaux': patient.get('ant_chirurgicaux'), 
                'ant_familiaux': patient.get('ant_familiaux'), 
                'traitements_actuels': patient.get('traitements_actuels')
            }
            
            for key, val in field_map.items():
                safe_set(key, val)
                
            # Defaults if empty
            if not self.widgets['wilaya'].get(): self.widgets['wilaya'].set(WILAYAS[0])
            self.update_communes(self.widgets['wilaya'].get())
            if not self.widgets['commune'].get(): self.widgets['commune'].set(patient.get('commune') or "")

            # Date
            try:
                date_str = patient.get('date_naissance')
                if date_str:
                    self.widgets['date_naissance'].set_date(datetime.strptime(date_str, "%d/%m/%Y"))
                else:
                    self.widgets['date_naissance'].delete(0, "end")
            except Exception as e: 
                print(f"Date Set Error: {e}")
            
            self.tuteur_lien_var.set(patient.get('tuteur_lien') or "Père")
            
            self.calculate_bmi()
            self.check_pediatric()
            
            self.load_interventions(pid)
            self.load_documents(pid)
        except Exception as e:
            print(f"Load Patient Error: {e}")
            messagebox.showerror("Erreur", f"Erreur chargement: {e}")

    def clear_form(self):
        self.current_patient = None
        for w in self.widgets.values():
            try: w.delete(0, "end")
            except: pass
            try: w.delete("1.0", "end")
            except: pass
            
        self.widgets['sexe'].set("Masculin")
        
        # Biskra Default
        biskra = next((w for w in WILAYAS if w.startswith("07")), WILAYAS[0])
        self.widgets['wilaya'].set(biskra)
        self.update_communes(biskra)
        
        self.widgets['groupe_sanguin'].set("A+")
        self.widgets['allergies'].set("Aucune")
        self.bmi_label.configure(text="IMC: --")
        self.age_label.configure(text="")
        self.pediatric_frame.pack_forget()
        
        for w in self.interv_list.winfo_children(): w.destroy()
        for w in self.doc_list.winfo_children(): w.destroy()

    def save_patient(self):
        import re
        phone = self.widgets['telephone'].get()
        # Allow spaces and optional leading 0. Format: usually 9 or 10 digits.
        clean_phone = phone.replace(" ", "").replace("-", "")
        if clean_phone and not re.match(r"^(0?)(5|6|7)[0-9]{8}$", clean_phone):
             messagebox.showerror("Erreur", "Format numéro de téléphone invalide (Ex: 550 35 23 33)")
             return

        dob_date = self.widgets['date_naissance'].get_date()
        data = {
            'nom': self.widgets['nom'].get(),
            'prenom': self.widgets['prenom'].get(),
            'nss': self.widgets['nss'].get(),
            'date_naissance': dob_date.strftime("%d/%m/%Y") if dob_date else "",
            'sexe': self.widgets['sexe'].get(),
            'telephone': self.widgets['telephone'].get(),
            'wilaya': self.widgets['wilaya'].get(),
            'commune': self.widgets['commune'].get(),
            'poids': self.widgets['poids'].get(),
            'taille': self.widgets['taille'].get(),
            'groupe_sanguin': self.widgets['groupe_sanguin'].get(),
            'allergies': self.widgets['allergies'].get(),
            # Fix: Textbox requires indices
            'vaccins': self.widgets['vaccins'].get("1.0", "end-1c"),
            'motif_consultation': self.widgets['motif_consultation'].get(),
            'ant_medicaux': self.widgets['ant_medicaux'].get("1.0", "end-1c"),
            'ant_chirurgicaux': self.widgets['ant_chirurgicaux'].get("1.0", "end-1c"),
            'ant_familiaux': self.widgets['ant_familiaux'].get("1.0", "end-1c"),
            'traitements_actuels': self.widgets['traitements_actuels'].get("1.0", "end-1c"),
            'tuteur_nom': self.widgets['tuteur_nom'].get(),
            'tuteur_lien': self.tuteur_lien_var.get(),
            # Pediatric
            'terme_grossesse': self.widgets['terme_grossesse'].get(),
            'poids_naissance': self.widgets['poids_naissance'].get(),
            'apgar': self.widgets['apgar'].get(),
            'mode_accouchement': self.widgets['mode_accouchement'].get(),
            # Legacy mapping
            'sang': self.widgets['groupe_sanguin'].get(),
            'parents': f"{self.tuteur_lien_var.get()}: {self.widgets['tuteur_nom'].get()}",
            'antecedents': "Voir détails"
        }
        
        if not data['nom'] or not data['prenom']:
            messagebox.showerror("Erreur", "Nom et Prénom obligatoires")
            return
            
        saved_id = None
        if self.current_patient:
            saved_id = self.current_patient['id']
            database.update_patient(saved_id, data)
        else:
            saved_id = database.add_patient(data)
            
        # Success Notification (Green Label instead of Popup)
        lbl_success = ctk.CTkLabel(self.btn_frame, text="✅ Dossier Sauvegardé !", text_color="green", font=("Roboto", 14, "bold"))
        lbl_success.pack(side="left", padx=20)
        # Auto-hide after 3 seconds
        self.after(3000, lambda: lbl_success.destroy())
            
        self.load_patients_list()
        
        # Auto-select the patient
        # Since add_patient doesn't return ID easily without modifying DB, 
        # I'll just rely on the user clicking, OR I can find by Name.
        # But user asked: "page should refresh and show the name i registreed on the side pannel"
        # It DOES show on the side panel after refresh.
        # "and click on it's name on the side i don't see his data" -> That was the bug. I fixed `load_patient`.
        # I will keep it simple: Refresh list is enough, unless they want auto-load.


    def delete_patient(self):
        if not self.current_patient: return
        if messagebox.askyesno("Confirmation", "Supprimer ce patient ?"):
            database.delete_patient(self.current_patient['id'])
            self.clear_form()
            self.load_patients_list()

    def export_patient_docx(self):
        if not self.current_patient: return
        
        file_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word Document", "*.docx")])
        if not file_path: return
        
        try:
            doc = docx.Document()
            p = self.current_patient
            
            # TITLE
            title = doc.add_heading(f"Dossier Médical - {p['nom']} {p['prenom']}", 0)
            title.alignment = 1 # Center
            
            # IDENTITY
            doc.add_heading("Identité & Contact", level=1)
            p_text = f"""
            NSS: {p.get('nss', '')}
            Date de Naissance: {p.get('date_naissance', '')} ({self.age_label.cget('text')})
            Sexe: {p.get('sexe', '')}
            Adresse: {p.get('commune', '')}, {p.get('wilaya', '')}
            Téléphone: {p.get('telephone', '')}
            Tuteur: {p.get('tuteur_lien', '')} - {p.get('tuteur_nom', '')}
            """
            doc.add_paragraph(p_text)
            
            # VITALS
            doc.add_heading("Données Cliniques", level=1)
            v_text = f"""
            Groupe Sanguin: {p.get('groupe_sanguin', '')}
            Poids: {p.get('poids', '')} kg   |   Taille: {p.get('taille', '')} cm   |   {self.bmi_label.cget('text')}
            Allergies: {p.get('allergies', '')}
            Vaccins: {p.get('vaccins', '').replace(chr(10), ', ')}
            Motif Consultation: {p.get('motif_consultation', '')}
            """
            doc.add_paragraph(v_text)
            
            if p.get('mode_accouchement'):
                doc.add_heading("Volet Pédiatrique / Naissance", level=2)
                ped_text = f"""
                Accouchement: {p.get('mode_accouchement', '')}
                Terme: {p.get('terme_grossesse', '')} SA
                Poids Naissance: {p.get('poids_naissance', '')} g
                APGAR: {p.get('apgar', '')}
                """
                doc.add_paragraph(ped_text)

            # HISTORY
            doc.add_heading("Antécédents", level=1)
            doc.add_paragraph(f"Médicaux:\n{p.get('ant_medicaux', '')}")
            doc.add_paragraph(f"Chirurgicaux:\n{p.get('ant_chirurgicaux', '')}")
            doc.add_paragraph(f"Familiaux:\n{p.get('ant_familiaux', '')}")
            doc.add_paragraph(f"Traitements en Cours:\n{p.get('traitements_actuels', '')}")

            # INTERVENTIONS
            doc.add_heading("Historique Interventions (DZSynapse)", level=1)
            interventions = database.get_interventions(p['id'])
            if interventions:
                table = doc.add_table(rows=1, cols=4)
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = 'Date'
                hdr_cells[1].text = 'Type'
                hdr_cells[2].text = 'Geste'
                hdr_cells[3].text = 'Diagnostic'
                
                for i in interventions:
                    row_cells = table.add_row().cells
                    row_cells[0].text = str(i['date'])
                    row_cells[1].text = str(i['type_chirurgie'])
                    row_cells[2].text = str(i['geste'])
                    row_cells[3].text = str(i['diagnostic'])
            else:
                doc.add_paragraph("Aucune intervention enregistrée.")

            doc.save(file_path)
            messagebox.showinfo("Export", "Document DOCX généré avec succès ✅")
            
        except Exception as e:
            messagebox.showerror("Erreur Export", f"Erreur: {e}")

    def delete_patient(self):
        if not self.current_patient: return
        if messagebox.askyesno("Confirmation", "Supprimer ce patient ?"):
            database.delete_patient(self.current_patient['id'])
            self.clear_form()
            self.load_patients_list()

    def load_interventions(self, pid):
        for w in self.interv_list.winfo_children(): w.destroy()
        try:
            interventions = database.get_interventions(pid)
            for i in interventions:
                card = ctk.CTkFrame(self.interv_list)
                card.pack(fill="x", pady=5)
                # Row access
                date_val = i['date']
                type_val = i['type_chirurgie']
                diag_val = i['diagnostic']
                geste_val = i['geste']
                
                bold_font = ctk.CTkFont(family="Roboto", size=12, weight="bold")
                ctk.CTkLabel(card, text=f"{date_val} - {type_val}", font=bold_font).pack(anchor="w", padx=5)
                ctk.CTkLabel(card, text=f"{diag_val} | {geste_val}", text_color="gray").pack(anchor="w", padx=5)
        except Exception as e:
             print(f"Error loading interventions: {e}")

    def load_documents(self, pid):
        for w in self.doc_list.winfo_children(): w.destroy()
        try:
            docs = database.get_documents(pid)
            for d_row in docs:
                d = dict(d_row) # Convert Row to Dict for .get() support
                # CARD CONTAINER
                card = ctk.CTkFrame(self.doc_list, fg_color="white", border_width=2, border_color="#e0e0e0", corner_radius=8)
                card.pack(fill="x", pady=5, padx=5)
                
                # 1. ICON PREVIEW (Simulated)
                # Determine icon based on type
                dtype = d['type'].lower()
                icon_text = "📄"
                icon_color = "#4CC9F0" # Default Blue
                if "image" in dtype or "radio" in dtype or "echo" in dtype:
                    icon_text = "🖼️"
                    icon_color = "#FFD700" # Gold
                elif "ordonnance" in dtype:
                    icon_text = "💊"
                    icon_color = "#00E676" # Green
                elif "biologique" in dtype:
                    icon_text = "🩸"
                    icon_color = "#FF5252" # Red

                icon_frame = ctk.CTkFrame(card, width=50, height=50, fg_color=icon_color, corner_radius=6)
                icon_frame.pack(side="left", padx=10, pady=10)
                icon_frame.pack_propagate(False) # Fixed size
                ctk.CTkLabel(icon_frame, text=icon_text, font=("Segoe UI Emoji", 24)).place(relx=0.5, rely=0.5, anchor="center")

                # 2. DETAILS
                info_frame = ctk.CTkFrame(card, fg_color="transparent")
                info_frame.pack(side="left", fill="both", expand=True, padx=10)
                
                ctk.CTkLabel(info_frame, text=d['type'], font=("Roboto", 14, "bold"), text_color="black", anchor="w").pack(fill="x")
                
                # Get Date (Handle missing for legacy)
                date_str = d.get('date_creation', 'Date inconnue')
                if not date_str: date_str = "Date inconnue"
                    
                ctk.CTkLabel(info_frame, text=f"Ajouté le: {date_str}", font=("Roboto", 11), text_color="gray50", anchor="w").pack(fill="x")
                
                # 3. ACTION
                path = d['file_path']
                btn = ctk.CTkButton(card, text="Ouvrir", width=80, **BUTTON_STYLE)
                btn.configure(command=lambda p=path: self.open_file(p))
                btn.pack(side="right", padx=15)
                
        except Exception as e:
            print(f"Error loading documents: {e}")

    def open_file(self, path):
        try:
            os.startfile(path)
        except Exception as e:
            messagebox.showerror("Erreur", f"Impossible d'ouvrir le fichier: {e}")

    def add_intervention_dialog(self):
        if not self.current_patient:
             messagebox.showwarning("Attention", "Veuillez d'abord sélectionner un patient.")
             return

        dialog = ctk.CTkToplevel(self)
        dialog.title("Nouvelle Intervention")
        dialog.geometry("450x650") # Taller for notes and spacing
        dialog.grab_set()
        
        # Scrollable container for the form
        container = ctk.CTkScrollableFrame(dialog)
        container.pack(fill="both", expand=True, padx=10, pady=10)
        
        entries = {}

        # 1. Date (CustomDatePicker)
        ctk.CTkLabel(container, text="Date:").pack(pady=(5, 0), anchor="w", padx=10)
        date_frame = ctk.CTkFrame(container, fg_color="transparent")
        date_frame.pack(fill="x", padx=10, pady=(0, 5))
        entries['date'] = CustomDatePicker(date_frame)
        entries['date'].pack(anchor="w")
        # Default to today
        entries['date'].set_date(datetime.now())

        # 2. Diagnostic (Entry)
        ctk.CTkLabel(container, text="Diagnostic:").pack(pady=(5, 0), anchor="w", padx=10)
        entries['diagnostic'] = ctk.CTkEntry(container)
        entries['diagnostic'].pack(fill="x", padx=10, pady=(0, 5))

        # 3. Type Chirurgie (ComboBox)
        ctk.CTkLabel(container, text="Type de Chirurgie:").pack(pady=(5, 0), anchor="w", padx=10)
        chirurgie_types = ["Chirurgie Générale", "Orthopédie", "Traumatologie", "Viscérale", "Urologie", "Gynécologie", "ORL", "Ophtalmologie", "Neurochirurgie", "Autre"]
        entries['type_chirurgie'] = ctk.CTkComboBox(container, values=chirurgie_types)
        entries['type_chirurgie'].pack(fill="x", padx=10, pady=(0, 5))

        # 4. Geste (Entry)
        ctk.CTkLabel(container, text="Geste Réalisé:").pack(pady=(5, 0), anchor="w", padx=10)
        entries['geste'] = ctk.CTkEntry(container)
        entries['geste'].pack(fill="x", padx=10, pady=(0, 5))

        # 5. Lateralite (OptionMenu - Bonus UX)
        ctk.CTkLabel(container, text="Latéralité:").pack(pady=(5, 0), anchor="w", padx=10)
        entries['lateralite'] = ctk.CTkOptionMenu(container, values=["Non Applicable", "Gauche", "Droite", "Bilatérale"])
        entries['lateralite'].pack(fill="x", padx=10, pady=(0, 5))

        # 6. Urgence (OptionMenu)
        ctk.CTkLabel(container, text="Urgence:").pack(pady=(5, 0), anchor="w", padx=10)
        entries['urgence'] = ctk.CTkOptionMenu(container, values=["Programmée (Froide)", "Urgence Relative", "Urgence Absolue", "Différée"])
        entries['urgence'].pack(fill="x", padx=10, pady=(0, 5))

        # 7. Notes (Textbox)
        ctk.CTkLabel(container, text="Notes Opératoires / Compléments:").pack(pady=(5, 0), anchor="w", padx=10)
        entries['notes'] = ctk.CTkTextbox(container, height=100)
        entries['notes'].pack(fill="x", padx=10, pady=(0, 5))

        def submit():
            # Get values safely
            try:
                date_obj = entries['date'].get_date()
                date_str = date_obj.strftime("%d/%m/%Y") if date_obj else ""
                
                data = {
                    'date': date_str,
                    'diagnostic': entries['diagnostic'].get(),
                    'type_chirurgie': entries['type_chirurgie'].get(),
                    'geste': entries['geste'].get(),
                    'lateralite': entries['lateralite'].get(),
                    'urgence': entries['urgence'].get(),
                    'notes': entries['notes'].get("1.0", "end-1c"),
                    'patient_id': self.current_patient['id']
                }
                
                database.add_intervention(data)
                self.load_interventions(self.current_patient['id'])
                dialog.destroy()
            except Exception as e:
                messagebox.showerror("Erreur", f"Erreur d'enregistrement: {e}")

        ctk.CTkButton(dialog, text="Ajouter Intervention", command=submit, **BUTTON_STYLE).pack(pady=10)

    def add_document_dialog(self):
        # ... keep existing ...
        if not self.current_patient:
             messagebox.showwarning("Attention", "Veuillez d'abord sélectionner un patient.")
             return
        file_path = tk.filedialog.askopenfilename()
        if file_path:
             dialog = ctk.CTkToplevel(self)
             dialog.title("Type de Document")
             dialog.geometry("300x200")
             dialog.grab_set()
             
             ctk.CTkLabel(dialog, text="Sélectionner le Type:").pack(pady=20)
             
             doc_types = ["Ordonnance", "Bilan Biologique", "Imagerie (Radio/Echo/IRM)", "Compte Rendu", "Certificat", "Lettre Orientation", "Autre"]
             type_entry = ctk.CTkComboBox(dialog, values=doc_types)
             type_entry.pack(padx=20, fill="x")
             
             def submit():
                 doc_type = type_entry.get() or "Autre"
                 storage_dir = os.path.join(os.getcwd(), "patient_docs", str(self.current_patient['id']))
                 os.makedirs(storage_dir, exist_ok=True)
                 filename = os.path.basename(file_path)
                 dest_path = os.path.join(storage_dir, filename)
                 try:
                     shutil.copy2(file_path, dest_path)
                     database.add_document({'patient_id': self.current_patient['id'], 'type': doc_type, 'file_path': dest_path})
                     self.load_documents(self.current_patient['id'])
                     dialog.destroy()
                 except Exception as e:
                     messagebox.showerror("Erreur", f"Échec: {e}")
             ctk.CTkButton(dialog, text="Enregistrer", command=submit, **BUTTON_STYLE).pack(pady=20)


class Sidebar(ctk.CTkFrame):
    def __init__(self, master, user, callbacks):
        # Set fg_color to white/light-gray to support black logo text
        super().__init__(master, width=260, corner_radius=0, fg_color=("white", "gray95"))
        self.callbacks = callbacks
        self.role = user['role']
        
        # LOGO INTEGRATION
        try:
            # Check for PyInstaller temp path first
            if hasattr(sys, '_MEIPASS'):
                 logo_path = os.path.join(sys._MEIPASS, "logo.png")
            else:
                 logo_path = os.path.join(os.getcwd(), "logo.png")
                 
            if os.path.exists(logo_path):
                # Load original image
                pil_image = Image.open(logo_path)
                self.logo_image = ctk.CTkImage(light_image=pil_image, dark_image=pil_image, size=(240, 240))
                
                # Transparent background now works because sidebar is white
                ctk.CTkLabel(self, image=self.logo_image, text="").pack(pady=(30, 10))
        except Exception as e:
            print(f"Logo load error: {e}")
        
        # Update text color to black for visibility on white
        # Greeting Update: "Bonjour, Docteur" (ignoring username for display)
        ctk.CTkLabel(self, text="Bonjour,\nDocteur", font=("Roboto", 18, "bold"), text_color="black").pack(pady=(0, 30))
        ctk.CTkLabel(self, text=f"Role: {self.role}", font=("Roboto", 12), text_color="gray20").pack(pady=(0, 20))
        
        # Buttons need to stand out on white. Primary blue is fine.
        self.btn_search = ctk.CTkButton(self, text="Dossier Patient", command=callbacks['patient'], **BUTTON_STYLE)
        self.btn_search.pack(pady=10, padx=20, fill="x")
        
        self.btn_kb = ctk.CTkButton(self, text="Référentiel Procédures", command=callbacks['reference'], **BUTTON_STYLE)
        self.btn_kb.pack(pady=10, padx=20, fill="x")
        
        ctk.CTkFrame(self, height=2, fg_color="gray80").pack(fill="x", pady=20, padx=20) # Lighter separator
        
        self.btn_logout = ctk.CTkButton(self, text="Déconnexion", command=callbacks['logout'], **BUTTON_STYLE_RED)
        self.btn_logout.pack(side="bottom", pady=20, padx=20, fill="x")
        
        # Theme Toggle (Sun/Moon Colors)
        self.mode_switch = ctk.CTkSwitch(self, text="Mode Sombre", 
                                         command=self.toggle_mode,
                                         progress_color="#183153",
                                         fg_color="#73C0FC",
                                         button_color="#e8e8e8",
                                         button_hover_color="#d0d0d0",
                                         text_color="black",
                                         font=("Roboto", 12, "bold"))
        self.mode_switch.pack(side="bottom", pady=(0, 10), padx=20, fill="x")
        self.mode_switch.select() # Default to Dark on load

        # FOOTER CREDITS
        # "Developed by DevOpsSolutions"
        # Link: https://devops-portfolio-bb6ee.web.app (Placeholder based on context)
        # Phone: +213 550 35 23 33 (Based on previous context)
        
        footer_frame = ctk.CTkFrame(self, fg_color="transparent")
        footer_frame.pack(side="bottom", pady=10)
        
        ctk.CTkLabel(footer_frame, text="Developed by", font=("Roboto", 10), text_color="gray50").pack()
        link = ctk.CTkLabel(footer_frame, text="DevOpsSolutions", font=("Roboto", 11, "bold", "underline"), 
                            text_color="#0b76ef", cursor="hand2")
        link.pack()
        link.bind("<Button-1>", lambda e: self.open_portfolio())
        
        ctk.CTkLabel(footer_frame, text="0772 38 37 79", font=("Roboto", 10, "bold"), text_color="gray40").pack()

    def open_portfolio(self):
        import webbrowser
        webbrowser.open("https://devops-solutions-2026.web.app/")

    def toggle_mode(self):
        if self.mode_switch.get() == 1:
            ctk.set_appearance_mode("Dark")
        else:
            ctk.set_appearance_mode("Light")

class App(ctk.CTk):
    def __init__(self):
        super().__init__()
        self.title("DZSynapse - Powered by DevOpsSolutions")
        self.geometry("1100x700")
        
        # Set App Icon
        try:
            logo_path = os.path.join(os.getcwd(), "logo.png")
            if os.path.exists(logo_path):
                # Using PIL to ImageTk for iconphoto
                icon_img = tk.PhotoImage(file=logo_path)
                self.iconphoto(False, icon_img)
        except Exception as e:
            print(f"Icon set error: {e}")
        
        self.login_frame = None
        self.main_view = None
        self.user = None
        
        # Initialize DB
        database.init_db()
        

        
        # Check for active session
        # Check for License first
        if not self.check_license():
            self.show_license()
        elif self.check_session():
            return
        else:
            self.show_login()

    def check_license(self):
        return os.path.exists("license.dat")

    def show_license(self):
        if self.main_view: self.main_view.destroy()
        self.login_frame = LicenseFrame(self, self.on_license_valid)

    def on_license_valid(self):
        # Create license file
        try:
            with open("license.dat", "w") as f:
                f.write("ACTIVATED")
        except:
            pass
            
        if self.login_frame:
            self.login_frame.destroy()
            
        if self.check_session():
            return
        self.show_login()

    def check_session(self):
        if os.path.exists("session.json"):
            try:
                with open("session.json", "r") as f:
                    data = json.load(f)
                user = database.validate_session(data.get("token"))
                if user and user['username'] == data.get("username"):
                    self.on_login_success(user)
                    return True
            except Exception as e:
                print(f"Session check failed: {e}")
        return False

    def show_login(self):
        if self.main_view:
            self.main_view.destroy()
        self.login_frame = LoginFrame(self, self.on_login_success)

    def on_login_success(self, user):
        self.user = user
        if self.login_frame:
            self.login_frame.destroy()
        self.setup_main_view()

    def setup_main_view(self):
        self.main_view = ctk.CTkFrame(self)
        self.main_view.pack(fill="both", expand=True)
        
        # Grid layout for Sidebar + Content
        self.main_view.grid_columnconfigure(1, weight=1)
        self.main_view.grid_rowconfigure(0, weight=1)
        
        self.content_frame = ctk.CTkFrame(self.main_view)
        self.content_frame.grid(row=0, column=1, sticky="nswe", padx=10, pady=10)
        
        callbacks = {
            'patient': lambda: self.show_frame("patient"),
            'reference': lambda: self.show_frame("reference"),
            'logout': self.logout
        }
        
        self.sidebar = Sidebar(self.main_view, self.user, callbacks)
        self.sidebar.grid(row=0, column=0, sticky="nswe")
        
        self.frames = {}
        
        # Initialize frames
        self.frames["patient"] = PatientManagementFrame(self.content_frame, self.user['role'])
        self.frames["reference"] = KnowledgeBaseFrame(self.content_frame)
        
        self.show_frame("patient")
        
    def show_frame(self, name):
        # Hide all frames
        for frame in self.frames.values():
            frame.pack_forget()
        # Show selected
        self.frames[name].pack(fill="both", expand=True)

    def logout(self):
        if self.user:
            database.clear_session(self.user['username'])
        if os.path.exists("session.json"):
            os.remove("session.json")
            
        self.user = None
        if self.main_view:
            self.main_view.destroy()
        self.show_login()

if __name__ == "__main__":
    app = App()
    app.mainloop()
